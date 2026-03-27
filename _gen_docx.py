"""Generate a clean OOXML-compliant DOCX from the PPA HTML spec."""
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# ── Parse the HTML into a simple AST ──────────────────────────────
class SpecParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.elements = []       # list of (tag, attrs_dict, text)
        self._stack = []         # open-tag stack
        self._buf = ""
        self._in_body = False
        self._skip = {"style", "head", "title"}
        self._cur_attrs = {}
        self._in_li = False
        self._ol_counter = 0
        self._list_type = None   # 'ol' or 'ul'

    def handle_starttag(self, tag, attrs):
        attrs_d = dict(attrs)
        if tag == "body":
            self._in_body = True
            return
        if not self._in_body or tag in self._skip:
            self._stack.append(tag)
            return
        if tag in ("h1", "h2", "h3", "p"):
            self._flush()
            self._stack.append(tag)
            self._cur_attrs = attrs_d
            self._buf = ""
        elif tag == "ol":
            self._flush()
            self._list_type = "ol"
            self._ol_counter = 0
        elif tag == "ul":
            self._flush()
            self._list_type = "ul"
        elif tag == "li":
            self._flush()
            self._in_li = True
            if self._list_type == "ol":
                self._ol_counter += 1
            self._stack.append("li")
            self._buf = ""
        elif tag in ("strong", "b", "em", "i"):
            self._stack.append(tag)
        else:
            self._stack.append(tag)

    def handle_endtag(self, tag):
        if tag == "body":
            self._flush()
            self._in_body = False
            return
        if not self._in_body:
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()
            return
        if tag in ("h1", "h2", "h3", "p"):
            self._flush()
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()
        elif tag == "li":
            self._flush()
            self._in_li = False
            if self._stack and self._stack[-1] == "li":
                self._stack.pop()
        elif tag in ("ol", "ul"):
            self._list_type = None
        elif tag in ("strong", "b", "em", "i"):
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()

    def handle_data(self, data):
        if not self._in_body:
            return
        self._buf += data

    def _flush(self):
        text = self._buf.strip()
        if not text:
            self._buf = ""
            return
        # Collapse whitespace
        text = re.sub(r"\s+", " ", text)
        # Determine element type
        cur_tag = None
        for t in reversed(self._stack):
            if t in ("h1", "h2", "h3", "p", "li"):
                cur_tag = t
                break
        if cur_tag is None:
            cur_tag = "p"  # default
        self.elements.append((cur_tag, dict(self._cur_attrs), text,
                              self._list_type, self._ol_counter))
        self._buf = ""


# ── Reparse with run-level bold/italic tracking ──────────────────
class RunParser(HTMLParser):
    """Second-pass parser that emits (tag, attrs, runs) where runs have bold/italic info."""
    def __init__(self):
        super().__init__()
        self.elements = []
        self._in_body = False
        self._skip = {"style", "head", "title"}
        self._tag_stack = []
        self._runs = []       # list of (text, bold, italic)
        self._cur_tag = None
        self._cur_attrs = {}
        self._bold = 0
        self._italic = 0
        self._list_type = None
        self._ol_counter = 0

    def handle_starttag(self, tag, attrs):
        attrs_d = dict(attrs)
        if tag == "body":
            self._in_body = True
            return
        if not self._in_body:
            return
        if tag in ("h1", "h2", "h3", "p"):
            self._flush_element()
            self._cur_tag = tag
            self._cur_attrs = attrs_d
            self._runs = []
        elif tag == "ol":
            self._flush_element()
            self._list_type = "ol"
            self._ol_counter = 0
        elif tag == "ul":
            self._flush_element()
            self._list_type = "ul"
        elif tag == "li":
            self._flush_element()
            self._cur_tag = "li"
            if self._list_type == "ol":
                self._ol_counter += 1
            self._runs = []
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1

    def handle_endtag(self, tag):
        if tag == "body":
            self._flush_element()
            self._in_body = False
            return
        if not self._in_body:
            return
        if tag in ("h1", "h2", "h3", "p", "li"):
            self._flush_element()
        elif tag in ("ol", "ul"):
            self._list_type = None
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)

    def handle_data(self, data):
        if not self._in_body:
            return
        if self._cur_tag:
            self._runs.append((data, self._bold > 0, self._italic > 0))

    def _flush_element(self):
        if self._cur_tag and self._runs:
            self.elements.append((
                self._cur_tag,
                dict(self._cur_attrs),
                list(self._runs),
                self._list_type,
                self._ol_counter,
            ))
        self._cur_tag = None
        self._runs = []
        self._cur_attrs = {}


# ── Build the DOCX ───────────────────────────────────────────────
def build_docx(html_path: str, out_path: str):
    with open(html_path) as f:
        html = f.read()

    parser = RunParser()
    parser.feed(html)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for tag, attrs, runs, list_type, ol_counter in parser.elements:
        if tag == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            full_text = "".join(r[0] for r in runs).strip()
            full_text = re.sub(r"\s+", " ", full_text)
            run = p.add_run(full_text.upper())
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif tag == "h2":
            p = doc.add_paragraph()
            full_text = "".join(r[0] for r in runs).strip()
            full_text = re.sub(r"\s+", " ", full_text)
            run = p.add_run(full_text.upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif tag == "h3":
            p = doc.add_paragraph()
            full_text = "".join(r[0] for r in runs).strip()
            full_text = re.sub(r"\s+", " ", full_text)
            run = p.add_run(full_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif tag == "p":
            is_center = attrs.get("class") == "center"
            p = doc.add_paragraph()
            if is_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for text, bold, italic in runs:
                cleaned = re.sub(r"\s+", " ", text)
                if not cleaned:
                    continue
                run = p.add_run(cleaned)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

        elif tag == "li":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            # Add bullet or number prefix
            if list_type == "ol":
                prefix_run = p.add_run(f"{ol_counter}. ")
                prefix_run.font.name = "Times New Roman"
                prefix_run.font.size = Pt(12)
            else:
                prefix_run = p.add_run("\u2022 ")
                prefix_run.font.name = "Times New Roman"
                prefix_run.font.size = Pt(12)
            for text, bold, italic in runs:
                cleaned = re.sub(r"\s+", " ", text)
                if not cleaned:
                    continue
                run = p.add_run(cleaned)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_docx(
        "InvertorsDocs/PPA-Specification-Draft.html",
        "InvertorsDocs/PPA-Specification-Draft.docx",
    )
